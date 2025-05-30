from flask import Flask, request, jsonify, send_from_directory, Response, stream_with_context
from flask_cors import CORS
import json
import requests
import os # 用于处理文件名和路径
from werkzeug.utils import secure_filename # 用于安全地处理上传的文件名
import pdfplumber  # 用于PDF解析
import docx        # 用于DOCX解析
import time

app = Flask(__name__, static_folder='.', static_url_path='') # 明确指定静态文件目录和URL路径
CORS(app)

# 确保上传文件夹存在
UPLOAD_FOLDER = './uploads' # 您可以指定服务器上的任何安全路径
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# DeepSeek API 配置 (安全警告：生产环境不要硬编码API Key)
DEEPSEEK_API_KEY = "sk-94ee67f1837740db9c8c9d0f2646b210"
DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"

@app.route('/')
def serve_index():
    return send_from_directory(app.static_folder, '职业规划小助手.html')

@app.route('/职业规划小助手.html')
def serve_html_directly():
    return send_from_directory(app.static_folder, '职业规划小助手.html')

# 辅助函数: 从DeepSeek API获取流式响应
def stream_deepseek_api(messages):
    response = requests.post(
        DEEPSEEK_API_URL,
        headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
        json={
            "model": "deepseek-chat",
            "messages": messages,
            "temperature": 0.7,  # 提高温度以加快生成速度
            "max_tokens": 2500,   # 控制生成长度
            "stream": True        # 启用流式输出
        },
        stream=True  # 使用requests的流式处理
    )
    
    if not response.ok:
        yield f"data: {json.dumps({'error': f'API错误: {response.status_code}'})}\n\n"
        return
    
    all_content = ""  # 用于累积完整内容
    
    for line in response.iter_lines():
        if line:
            line_text = line.decode('utf-8')
            
            # DeepSeek API流式响应格式处理
            if line_text.startswith('data: '):
                data_str = line_text[6:]  # 去掉 'data: ' 前缀
                
                if data_str == "[DONE]":
                    # 流结束，发送完整内容
                    yield f"data: {json.dumps({'done': True, 'full_content': all_content})}\n\n"
                    break
                
                try:
                    data = json.loads(data_str)
                    if 'choices' in data and len(data['choices']) > 0:
                        delta = data['choices'][0].get('delta', {})
                        content = delta.get('content', '')
                        if content:
                            all_content += content
                            yield f"data: {json.dumps({'content': content, 'done': False})}\n\n"
                except Exception as e:
                    print(f"解析流式响应时出错: {e}")
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"

# 辅助函数: 清理JSON字符串
def clean_json_string(json_str):
    # 去除可能包含的Markdown代码块符号
    if json_str.startswith('```') and json_str.endswith('```'):
        json_str = json_str.strip('`')
        
    # 尝试找到JSON块的开始和结束
    start_idx = json_str.find('{')
    end_idx = json_str.rfind('}')
    
    if start_idx != -1 and end_idx != -1:
        return json_str[start_idx:end_idx+1]
    return json_str

# 提取简历内容
def extract_resume_content(file):
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    text = ""
    try:
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.pdf':
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    
        elif file_ext == '.docx':
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
                
        else:
            text = f"不支持的文件格式: {file_ext}"
            
    except Exception as e:
        text = f"解析文件时出错: {str(e)}"
        
    return text

# 流式职业分析API
@app.route('/api/analyze/stream', methods=['POST'])
def stream_analyze_career():
    mbti = request.form.get('mbti', '未提供')
    city = request.form.get('city', '未提供')
    resume_content = "未上传简历"
    
    # 解析霍兰德测试答案
    holland_answers_str = request.form.get('holland_answers', '{}')
    try:
        holland_answers = json.loads(holland_answers_str)
    except:
        holland_answers = {}
    
    # 简化的霍兰德代码解析
    holland_code = {}
    for q, answer in holland_answers.items():
        if answer not in holland_code:
            holland_code[answer] = 0
        holland_code[answer] = holland_code[answer] + 1
    
    holland_summary = ", ".join([f"{code}:{count}" for code, count in holland_code.items()])
    
    # 处理上传的简历
    resume_file = request.files.get('resume')
    if resume_file and resume_file.filename:
        resume_content = extract_resume_content(resume_file)
        # 简化简历内容以加快API响应
        if len(resume_content) > 2000:
            resume_content = resume_content[:2000] + "...(已截断)"
    
    # 构建提示消息
    system_message = {
        "role": "system",
        "content": """你是专业的职业规划分析师。请基于用户的MBTI、霍兰德测试结果和简历内容进行分析。
你的任务是：
1. 提供整体分析和洞察(insight)，包括三个明确部分：
   - MBTI分析：分析MBTI类型与职业特质的关系
   - 霍兰德分析：分析霍兰德模型结果与职业兴趣的关系
   - 简历分析：分析简历显示的能力、技能和背景
2. 推荐5个最匹配的职业(recommendations)，每个包含ID、名称、简短描述和推荐理由。

请直接返回分析结果，不需要加入markdown标记。结果应包含以下内容：
- insight: 整体分析，明确分为MBTI分析、霍兰德分析和简历分析三部分
- recommendations: 包含5个职业推荐，每个有id、name、short_description和reason字段

保持分析精确但简洁，注重质量和实用性。"""
    }
    
    user_message = {
        "role": "user",
        "content": f"""请分析以下信息并给出职业规划建议：
- MBTI类型：{mbti}
- 所在城市：{city}
- 霍兰德测试结果：{holland_summary}
- 简历内容：{resume_content}

请注意，您的回复将被直接解析为JSON格式，因此请确保您的回复是有效的JSON内容。
"""
    }
    
    messages = [system_message, user_message]
    
    # 使用stream_with_context确保流式响应可以正确传输
    return Response(stream_with_context(stream_deepseek_api(messages)), 
                   mimetype='text/event-stream')

# 流式职业详情API
@app.route('/api/career-details/stream', methods=['POST'])
def stream_career_details():
    career_id = request.form.get('career_id', '')
    career_name = request.form.get('career_name', '未知职业')
    city = request.form.get('city', '未提供')
    mbti = request.form.get('mbti', '未提供')
    
    # 解析霍兰德测试结果和简历内容，与analyze_career类似
    holland_answers_str = request.form.get('holland_answers', '{}')
    try:
        holland_answers = json.loads(holland_answers_str)
    except:
        holland_answers = {}
    
    holland_code = {}
    for q, answer in holland_answers.items():
        if answer not in holland_code:
            holland_code[answer] = 0
        holland_code[answer] = holland_code[answer] + 1
    
    holland_summary = ", ".join([f"{code}:{count}" for code, count in holland_code.items()])
    
    resume_content = "未提供简历"
    resume_file = request.files.get('resume')
    if resume_file and resume_file.filename:
        resume_content = extract_resume_content(resume_file)
        if len(resume_content) > 1000:  # 简化简历内容
            resume_content = resume_content[:1000] + "...(已截断)"
    
    system_message = {
        "role": "system",
        "content": """你是专业的职业规划分析师，现在需要针对用户选择的特定职业提供详细分析。
根据用户的背景（MBTI、霍兰德测试、简历）以及所选职业，请提供以下详细内容：
1. name: 职业名称
2. salary_info: 该职业在用户所在城市的薪资范围和影响因素
3. development_plan: 该职业的发展路径和晋升规划
4. learning_resources: 进入和提升该职业所需的学习资源和证书
5. core_competencies: 该职业所需的核心能力和素质
6. daily_workflow: 该职业的日常工作内容和节奏

请直接以JSON格式返回上述内容，确保每个字段都有具体且有用的信息。不要返回没有内容的字段。"""
    }
    
    user_message = {
        "role": "user",
        "content": f"""请为以下职业提供详细分析：
- 职业：{career_name} (ID: {career_id})
- MBTI类型：{mbti}
- 所在城市：{city}
- 霍兰德测试结果：{holland_summary}
- 简历内容：{resume_content}

请以JSON格式直接返回结果，包含name、salary_info、development_plan、learning_resources、core_competencies和daily_workflow字段。"""
    }
    
    messages = [system_message, user_message]
    
    return Response(stream_with_context(stream_deepseek_api(messages)),
                   mimetype='text/event-stream')

# --- 新增非流式API调用辅助函数 ---
def call_deepseek_api_once(messages):
    """
    调用DeepSeek API并一次性获取完整响应。
    确保 'stream': False。
    """
    response = requests.post(
        DEEPSEEK_API_URL,
        headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
        json={
            "model": "deepseek-chat",
            "messages": messages,
            "temperature": 0.7,
            "max_tokens": 4000, # 调整Token以适应可能更长的内容
            "stream": False # 重要：确保为非流式调用
        }
    )
    if response.ok:
        try:
            response_data = response.json()
            raw_response_content = response_data['choices'][0]['message']['content']
            # 使用已有的清理函数
            cleaned_json_str = clean_json_string(raw_response_content)
            try:
                return json.loads(cleaned_json_str)
            except json.JSONDecodeError as e:
                error_message = f"JSON解析错误: {str(e)}. AI原始返回 (已清理): {cleaned_json_str}"
                print(error_message)
                return {"error": "AI响应格式错误，无法解析为JSON", "details": error_message, "raw_content_preview": cleaned_json_str[:500]}
        except (KeyError, IndexError) as e:
            error_message = f"解析API响应结构时出错: {str(e)}. 完整响应: {response.text[:500]}"
            print(error_message)
            return {"error": "AI响应结构意外", "details": error_message}
        except Exception as e: # 其他可能的请求或解析错误
            error_message = f"处理API响应时发生未知错误: {str(e)}. 完整响应: {response.text[:500]}"
            print(error_message)
            return {"error": "处理AI响应时发生未知错误", "details": error_message}
    else:
        error_message = f"API请求失败: {response.status_code}. 响应: {response.text[:500]}"
        print(error_message)
        return {"error": f"AI服务API请求失败", "status_code": response.status_code, "details": error_message}

# --- 新的API端点：初步分析 (非流式) ---
@app.route('/api/analyze-preliminary', methods=['POST'])
def analyze_preliminary():
    mbti = request.form.get('mbti', '未提供')
    city = request.form.get('city', '未提供')
    holland_answers_str = request.form.get('holland_answers', '{}')
    try:
        holland_answers = json.loads(holland_answers_str)
    except json.JSONDecodeError:
        holland_answers = {}
    
    holland_code = {}
    for _q, answer in holland_answers.items(): # _q is not used
        if answer not in holland_code: holland_code[answer] = 0
        holland_code[answer] += 1
    holland_summary = ", ".join([f"{code}:{count}" for code, count in holland_code.items()])

    # 初步分析不包含简历内容
    resume_content_for_prompt = "用户未提供简历，请仅基于MBTI和霍兰德信息进行分析。"

    system_message_preliminary = {
        "role": "system",
        "content": """你是专业的职业规划分析师。请基于用户的MBTI和霍兰德测试结果进行初步分析。简历未提供。
你的任务是：
1. 提供整体分析和洞察(insight)。理想情况下，`insight`字段应为一个JSON对象，包含以下键：
   - "mbti_insight": 对MBTI类型的分析及其与职业特质的关系。
   - "holland_insight": 对霍兰德模型结果的分析及其与职业兴趣的关系。
   - "resume_insight": 对于此初步分析，请明确指出"用户未提供简历，此部分分析略过。"
2. 推荐3-5个匹配的职业(recommendations)，每个职业信息应包含id, name, short_description, 和 reason。这些推荐应主要基于MBTI和霍兰德信息。

返回结果必须是结构化的JSON，最外层包含 "insight" 和 "recommendations" 字段。
例如: {"insight": {"mbti_insight": "...", "holland_insight": "...", "resume_insight": "用户未提供简历，此部分分析略过。"}, "recommendations": [{"id":"1", "name":"职业A", "short_description":"...", "reason":"..."}, ...]}
如果无法生成特定部分的洞察，请在该部分说明原因，但保持整体JSON结构完整。确保所有文本内容都是UTF-8编码。不要在JSON字符串值中使用未转义的换行符，请使用 \\n。"""
    }
    
    user_message_preliminary = {
        "role": "user",
        "content": f"""请分析以下信息并给出初步职业规划建议（不含简历分析）：
- MBTI类型：{mbti}
- 所在城市：{city}
- 霍兰德测试结果：{holland_summary}
- 简历内容：{resume_content_for_prompt}

请严格按照指定的JSON结构返回结果。"""
    }
    
    messages = [system_message_preliminary, user_message_preliminary]
    analysis_result = call_deepseek_api_once(messages)
    return jsonify(analysis_result)

# --- 新的API端点：完整分析 (非流式, 包含简历) ---
@app.route('/api/analyze-full', methods=['POST'])
def analyze_full():
    mbti = request.form.get('mbti', '未提供')
    city = request.form.get('city', '未提供')
    holland_answers_str = request.form.get('holland_answers', '{}')
    try:
        holland_answers = json.loads(holland_answers_str)
    except json.JSONDecodeError:
        holland_answers = {}

    holland_code = {}
    for _q, answer in holland_answers.items(): # _q is not used
        if answer not in holland_code: holland_code[answer] = 0
        holland_code[answer] += 1
    holland_summary = ", ".join([f"{code}:{count}" for code, count in holland_code.items()])

    resume_content_for_prompt = "用户未上传简历。"
    resume_file = request.files.get('resume')
    if resume_file and resume_file.filename:
        try:
            extracted_text = extract_resume_content(resume_file) # 您已有的函数
            if len(extracted_text) > 3000: # 调整截断长度以提供更多上下文
                resume_content_for_prompt = extracted_text[:3000] + "...(简历内容过长，已截断)"
            else:
                resume_content_for_prompt = extracted_text
        except Exception as e:
            print(f"提取简历内容时出错: {e}")
            resume_content_for_prompt = "处理简历文件时发生错误，无法提取内容。"
    
    system_message_full = {
        "role": "system",
        "content": """你是专业的职业规划分析师。请基于用户的MBTI、霍兰德测试结果和简历内容（如果提供）进行全面分析。
你的任务是：
1. 提供整体分析和洞察(insight)。理想情况下，`insight`字段应为一个JSON对象，包含以下键：
   - "mbti_insight": 对MBTI类型的分析及其与职业特质的关系。
   - "holland_insight": 对霍兰德模型结果的分析及其与职业兴趣的关系。
   - "resume_insight": 对简历中显示的能力、经验和技能的分析。如果用户未提供简历，或简历内容无法有效分析，请明确说明。
2. 推荐5个最匹配的职业(recommendations)，每个职业信息应包含id, name, short_description, 和 reason。这些推荐应综合所有输入信息。

返回结果必须是结构化的JSON，最外层包含 "insight" 和 "recommendations" 字段。
例如: {"insight": {"mbti_insight": "...", "holland_insight": "...", "resume_insight": "..."}, "recommendations": [{"id":"1", "name":"职业A", "short_description":"...", "reason":"..."}, ...]}
确保所有文本内容都是UTF-8编码。不要在JSON字符串值中使用未转义的换行符，请使用 \\n。"""
    }
    
    user_message_full = {
        "role": "user",
        "content": f"""请分析以下信息并给出全面的职业规划建议：
- MBTI类型：{mbti}
- 所在城市：{city}
- 霍兰德测试结果：{holland_summary}
- 简历内容：{resume_content_for_prompt}

请严格按照指定的JSON结构返回结果。"""
    }
    
    messages = [system_message_full, user_message_full]
    analysis_result = call_deepseek_api_once(messages)
    return jsonify(analysis_result)

# 保留原有的非流式API端点，以兼容旧版本
@app.route('/api/analyze', methods=['POST'])
def analyze_career():
    mbti = request.form.get('mbti', '未提供')
    city = request.form.get('city', '未提供')
    holland_answers_str = request.form.get('holland_answers', '{}')
    
    try:
        holland_answers_dict = json.loads(holland_answers_str)
    except json.JSONDecodeError:
        holland_answers_dict = {}
        print(f"无法解析霍兰德答案字符串: {holland_answers_str}")

    resume_file = request.files.get('resume')
    resume_content_summary = "未上传简历。"
    resume_text = ""
    original_filename = ""

    # 解析简历内容
    if resume_file and resume_file.filename:
        original_filename = secure_filename(resume_file.filename)
        ext = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else ''
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], original_filename)
        resume_file.save(file_path)
        
        try:
            if ext == 'pdf':
                # 使用pdfplumber提取PDF文本
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            resume_text += extracted_text + "\n\n"
            elif ext in ['doc', 'docx']:
                # 使用python-docx提取Word文本
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    resume_text += para.text + "\n"
            else:
                resume_text = f"不支持的文件格式: {ext}"
                
            # 如果提取到内容，使用前1000个字符作为摘要
            if resume_text:
                resume_content_summary = resume_text[:1000] + ("..." if len(resume_text) > 1000 else "")
        except Exception as e:
            resume_content_summary = f"解析文件时出错: {str(e)}"

    # 简化的霍兰德测试结果分析 (R-现实型, I-研究型, A-艺术型, S-社会型, E-企业型, C-常规型)
    holland_codes = {'R': 0, 'I': 0, 'A': 0, 'S': 0, 'E': 0, 'C': 0}
    for q, answer in holland_answers_dict.items():
        if answer in holland_codes:
            holland_codes[answer] += 1
    
    # 找出得分最高的三种类型
    top_holland = sorted(holland_codes.items(), key=lambda x: x[1], reverse=True)[:3]
    holland_summary = ", ".join([f"{code}型({score}分)" for code, score in top_holland if score > 0])
    
    # 构建和AI模型的对话消息
    system_message = {
        "role": "system",
        "content": """你是专业的职业规划分析师。请基于用户的MBTI、霍兰德测试结果和简历内容进行分析。
你的回答必须是JSON格式，包含以下两部分：
1. insight: 对用户的整体分析，包括MBTI分析、霍兰德模型分析和简历分析
2. recommendations: 推荐5个职业，每个职业必须包含以下字段：
   - id: 唯一标识符（数字）
   - name: 职业名称
   - short_description: 简短描述
   - reason: 推荐原因
   - percentage: 匹配度百分比（1-100的数字）

确保按照上述格式返回，不要添加任何其他格式如markdown。只返回纯JSON格式内容。"""
    }
    
    user_message = {
        "role": "user",
        "content": f"""请分析以下信息并给出职业规划建议：
- MBTI类型：{mbti}
- 所在城市：{city}
- 霍兰德测试结果：{holland_summary}
- 简历内容：{resume_content_summary}

请注意，您的回复将被直接解析为JSON格式，因此请确保您的回复是有效的JSON，并且不要使用markdown格式(如```json)。
"""
    }
    
    # 调用DeepSeek API
    try:
        response = requests.post(
            DEEPSEEK_API_URL,
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            json={
                "model": "deepseek-chat",
                "messages": [system_message, user_message],
                "temperature": 0.7,
                "max_tokens": 3000
            }
        )
        
        if response.status_code != 200:
            return jsonify({"error": f"API调用失败: {response.status_code}", "detail": response.text}), 500
        
        api_response = response.json()
        if "choices" not in api_response or len(api_response["choices"]) == 0:
            return jsonify({"error": "API返回无效响应", "detail": api_response}), 500
        
        ai_response = api_response["choices"][0]["message"]["content"]
        
        # 清理并解析JSON响应
        try:
            cleaned_json = clean_json_string(ai_response)
            result = json.loads(cleaned_json)
            
            # 验证结果格式
            if not isinstance(result, dict):
                return jsonify({"error": "AI分析结果格式错误", "short_description": "返回的不是一个对象"}), 500
            
            if "insight" not in result or "recommendations" not in result:
                return jsonify({"error": "AI分析结果格式错误", "short_description": "缺少必要的字段"}), 500
            
            if not isinstance(result["recommendations"], list):
                return jsonify({"error": "AI分析结果格式错误", "short_description": "recommendations不是数组"}), 500
            
            # 验证职业推荐数据完整性
            for rec in result["recommendations"]:
                if not all(k in rec for k in ["id", "name", "short_description", "reason"]):
                    return jsonify({"error": "AI分析结果格式错误", "short_description": "AI返回的部分职业数据不完整。"}), 500
            
            return jsonify(result)
            
        except json.JSONDecodeError as e:
            print(f"JSON解析错误: {e}")
            print(f"原始响应: {ai_response}")
            return jsonify({"error": "AI分析响应非JSON", "short_description": str(e)}), 500
            
    except Exception as e:
        print(f"API调用出错: {str(e)}")
        return jsonify({"error": "调用AI分析失败", "short_description": str(e)}), 500

@app.route('/api/career-details', methods=['POST'])
def get_career_details():
    career_id = request.form.get('career_id', '')
    career_name = request.form.get('career_name', '未知职业')
    city = request.form.get('city', '未提供')
    mbti = request.form.get('mbti', '未提供')
    
    holland_answers_str = request.form.get('holland_answers', '{}')
    try:
        holland_answers_dict = json.loads(holland_answers_str)
    except json.JSONDecodeError:
        holland_answers_dict = {}
    
    # 处理简历文件
    resume_file = request.files.get('resume')
    resume_content = "未上传简历"
    if resume_file and resume_file.filename:
        filename = secure_filename(resume_file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        resume_file.save(file_path)
        
        try:
            # 根据文件扩展名处理
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            if ext == 'pdf':
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            resume_content += extracted_text + "\n\n"
            elif ext in ['doc', 'docx']:
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    resume_content += para.text + "\n"
            else:
                resume_content = f"不支持的文件格式: {ext}"
        except Exception as e:
            resume_content = f"解析文件时出错: {str(e)}"
    
    # 构建霍兰德测试结果摘要
    holland_codes = {'R': 0, 'I': 0, 'A': 0, 'S': 0, 'E': 0, 'C': 0}
    for q, answer in holland_answers_dict.items():
        if answer in holland_codes:
            holland_codes[answer] += 1
    
    top_holland = sorted(holland_codes.items(), key=lambda x: x[1], reverse=True)[:3]
    holland_summary = ", ".join([f"{code}型({score}分)" for code, score in top_holland if score > 0])
    
    # 构建提示消息
    system_message = {
        "role": "system",
        "content": """你是专业的职业规划分析师，现在需要针对用户选择的特定职业提供详细分析。
根据用户的背景（MBTI、霍兰德测试、简历）以及所选职业，请以JSON格式返回以下详细内容。确保所有字段都包含具体且有用的信息。

返回的JSON对象应包含以下顶级键，且**仅包含这些键**：
1.  `name`: (string) 职业名称。
2.  `salary_info`: (string) 该职业在用户所在城市的典型薪资范围、可能的变动因素（如经验、公司规模等）。
3.  `development_plan`: (string) 该职业的详细发展路径和晋升规划。请将其描述为分阶段的计划（例如：短期1-2年，中期3-5年，长期5年以上的目标、学习和成就）。使用清晰的段落或列表呈现每个阶段。
4.  `learning_resources`: (string) 进入和提升该职业所需的学习资源、推荐课程、在线平台、重要书籍以及建议考取的专业证书。
5.  `hard_skills`: (string) 从事该职业所需的具体硬技能列表或详细描述（例如：编程语言、软件工具、专业知识领域）。**此字段必须提供。**
6.  `soft_skills`: (string) 从事该职业所需的关键软技能列表或详细描述（例如：沟通能力、团队合作、解决问题、领导力）。**此字段必须提供。**
7.  `mbti_advantage`: (string) 结合用户提供的MBTI类型，分析该性格类型在该职业中的主要优势以及如何有效发挥这些优势的具体建议。**此字段必须提供。**
8.  `daily_workflow`: (string) 该职业典型的一天或一周工作内容。请以任务列表的形式描述（例如，用换行符分隔每个任务或活动），以便前端可以解析为日程表。

**重要：不要返回名为 `core_competencies` 的字段。核心能力必须分解到上述的 `hard_skills`、`soft_skills` 和 `mbti_advantage` 三个独立字段中。**

确保返回的JSON格式有效且可以直接使用，不要包含任何markdown代码块 (如\`\`\`json)。所有文本内容应该是中文。"""
    }
    
    user_message = {
        "role": "user",
        "content": f"""请为以下职业提供详细分析：
- 职业：{career_name} (ID: {career_id})
- MBTI类型：{mbti}
- 所在城市：{city}
- 霍兰德测试结果：{holland_summary}
- 简历内容：{resume_content}"""
    }
    
    # 调用DeepSeek API
    try:
        response = requests.post(
            DEEPSEEK_API_URL,
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            json={
                "model": "deepseek-chat",
                "messages": [system_message, user_message],
                "temperature": 0.7,
                "max_tokens": 3000
            }
        )
        
        if response.status_code != 200:
            return jsonify({"error": f"API调用失败: {response.status_code}"}), 500
        
        api_response = response.json()
        if "choices" not in api_response or len(api_response["choices"]) == 0:
            return jsonify({"error": "API返回无效响应"}), 500
        
        ai_response = api_response["choices"][0]["message"]["content"]
        
        # 清理并解析JSON响应
        try:
            cleaned_json = clean_json_string(ai_response)
            result = json.loads(cleaned_json)
            
            # 确保基本字段存在，如果不存在则添加默认值
            if "name" not in result:
                result["name"] = career_name
                
            return jsonify(result)
            
        except json.JSONDecodeError as e:
            print(f"JSON解析错误: {e}")
            print(f"原始响应: {ai_response}")
            return jsonify({
                "error": "API响应解析失败", 
                "name": career_name,
                "salary_info": "无法获取薪资信息",
                "development_plan": "<p>获取职业发展路径信息失败</p>",
                "learning_resources": "<p>获取学习资源信息失败</p>",
                "hard_skills": "<p>获取硬技能信息失败</p>",
                "soft_skills": "<p>获取软技能信息失败</p>",
                "mbti_advantage": "<p>获取MBTI优势分析失败</p>",
                "daily_workflow": "<p>获取日常工作流程信息失败</p>"
            })
            
    except Exception as e:
        print(f"API调用出错: {str(e)}")
        return jsonify({
            "error": str(e), 
            "name": career_name,
            "salary_info": "调用AI分析时出错",
            "development_plan": f"<p>错误: {str(e)}</p>",
            "learning_resources": "<p>请稍后再试</p>",
            "hard_skills": "<p>请稍后再试</p>",
            "soft_skills": "<p>请稍后再试</p>",
            "mbti_advantage": "<p>请稍后再试</p>",
            "daily_workflow": "<p>请稍后再试</p>"
        })

if __name__ == '__main__':
    # 确保在 /www/wwwroot/wm/ 目录下运行此脚本
    # 或者用绝对路径指定日志文件位置
    # logging.basicConfig(filename='app.log', level=logging.DEBUG) # 可以取消注释来记录日志到文件
    app.run(host='0.0.0.0', port=8000) 